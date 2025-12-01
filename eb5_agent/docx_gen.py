# Purpose: Generate Flow of Funds DOCX with sources, narrative, and investment statement.
from docx import Document
from datetime import date

def generate_flow_doc(case_id: str, client_name: str, project_name: str, sources: list, transactions: list) -> str:
    """Build a Flow of Funds DOCX file for a case."""
    doc = Document()
    doc.add_heading(f"Flow of Funds – {client_name}", level=1)
    doc.add_paragraph(f"Case ID: {case_id}")
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(f"Generated: {date.today().strftime('%B %d, %Y')}")

    # Sources section with optional date ranges
    doc.add_heading("Sources of Funds", level=2)
    for s in sources:
        amt = s.get("target_amount")
        dr = s.get("date_range")
        line = f"- {s['label']} ({s['source_type']})"
        if amt: line += f": ${amt:,}"
        if dr: line += f" accumulated between {dr[0]} and {dr[1]}"
        doc.add_paragraph(line + ".")

    # Flow narrative: include dates, amounts, suffixes where available, and addresses
    doc.add_heading("Flow Narrative", level=2)
    for t in transactions:
        origin_suffix = t.get("origin_suffix")
        dest_suffix = t.get("destination_suffix")
        line = (
            f"On {t['date']}, ${t['amount']:,} was transferred from {t['origin']}"
            + (f" (Acct ****{origin_suffix})" if origin_suffix else "")
            + f" to {t['destination']}"
            + (f" (Acct ****{dest_suffix})" if dest_suffix else "")
            + "."
        )
        if t.get("address"):  # institution/destination address if available
            line += f" Address: {t['address']}."
        doc.add_paragraph(line)

    # Investment statement (simple total; adapt to your template)
    total = sum([t['amount'] for t in transactions])
    doc.add_heading("Investment Statement", level=2)
    doc.add_paragraph(f"The investor is investing a total of ${total:,} into {project_name} as part of the EB-5 program.")

    filename = f"{case_id}_flow_of_funds.docx"
    doc.save(filename)
    return filename